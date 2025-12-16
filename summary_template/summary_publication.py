import os
import ast
import json
import random
import pandas as pd
import pandas as pd
from backend.scripts.utils import query_gai_via_gateway
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from nltk.corpus import stopwords
from nltk.stem import PorterStemmer
from nltk.tokenize import word_tokenize
import nltk
import datetime
from nltk.tokenize import sent_tokenize
from backend.scripts.utils import gai
# Download NLTK resources
nltk.download('punkt')
nltk.download('stopwords')
STOPWORDS = set(stopwords.words('english'))
STEMMER = PorterStemmer()


os.environ["DB_HOST"] = "localhost"
os.environ["POSTGRES_USER"] = "matthew50"
os.environ["POSTGRES_PASSWORD"] = "softpower"
os.environ["POSTGRES_DB"] = "softpower-db"
from backend.scripts.utils import Config
from backend.app import app
from backend.extensions import db
from backend.scripts.models import(SoftPowerActivity, 
                                   Document, 
                                   Category, 
                                   InitiatingCountry, 
                                   RecipientCountry, 
                                   DailySummary,
                                   DailyCategory,
                                   DailyRecipient,
                                   CountrySummary,
                                   Subcategory,
                                   Event,
                                   EventSources,
                                   EventSummary,
                                   TokenizedDocuments,
                                   SummarySources)
from dotenv import load_dotenv
from sqlalchemy.dialects.postgresql import insert
from sqlalchemy import func
load_dotenv()
cfg = Config.from_yaml()
app.app_context().push()
db.session.remove()
db.engine.dispose()
# reporting = (db.session.query(Document.doc_id,
#                             Document.date,
#                             Document.title,
#                             Document.distilled_text,
#                             Document.salience_justification,
#                             Category.category)
#                             .join(Category,Document.doc_id==Category.doc_id)
#                             .join(InitiatingCountry,Document.doc_id==InitiatingCountry.doc_id)
#                             .join(RecipientCountry,Document.doc_id==RecipientCountry.doc_id)
#                             .filter(Document.date.between(start,end),
#                             InitiatingCountry.initiating_country==country,
#                             RecipientCountry.recipient_country.in_(cfg.recipients),
#                             Category.category==category).all()
# )

def file_name(country,start_date,end_date,category,recipient):
    filename = ''
    if country:
        filename += f'{country}'
    if recipient:
        filename += f'_{recipient}'
    filename += f'_{start_date}_{end_date}'
    if category:
        filename += f'_{category}'
    return filename

# Preprocessing
def preprocess(text):
    tokens = word_tokenize(text.lower())
    filtered_tokens = [word for word in tokens if word not in STOPWORDS]
    stemmed_tokens = [STEMMER.stem(word) for word in filtered_tokens]
    # Open a session within the thread context
    return ' '.join(stemmed_tokens)

def find_related_articles(event_description, article_vectors, vectorizer, articles):
    # Vectorize the event description
    event_vector = vectorizer.transform([event_description])
    # Compute similarity
    similarities = cosine_similarity(event_vector, article_vectors)
    return similarities.flatten()

def source_summary(records,summary,summ_type):
    if summ_type=='overview':
        summ = summary['overview']
    else:
        summ = summary['outcome']
    sys_prompt = source_sys_prompt
    user_prompt = source_user_prompt.format(title=summary['key_event'], summary=summ, reports=records)
    source_response = gai(sys_prompt=sys_prompt, user_prompt=user_prompt,model="gpt-41")
    # source_output = fetch_gai_content(source_response)
    return source_response

def source_summaries(summaries, records):
    source_outputs = []
    if len(records) > 200:
        print('sampling records...')
        records = records.nlargest(200, 'similarity_score')
    # Source sentences for each summary
    for summary in summaries:
        print(f'processing {summary["id"]}')
        overview_output = source_summary(records,summary,summ_type='overview')
        outcome_output = source_summary(records,summary,summ_type='outcome')
        # insert_sentence_sourcing(gai_output=source_output,sentence=sentence,summary=summary)
    return overview_output,outcome_output

def fetch_gai_sources(output):
    sources = output['choices'][0]['message']['content'].replace('[','').replace(']','').replace("\n",'').replace("{","").replace("}","").replace("doc_id","").replace(":","").replace(' ','').strip().split(',')
    if sources:
        return sources 
    else:
        print(f"error: {sources}")
        return []
def clean_string(x):
    # Remove brackets, braces, parentheses, quotes, the word doc_id, colons, spaces, and newlines
    return re.sub(r"[\[\]\{\}\(\)'\:\s\n]|doc_id", "", x)

    def summary_reporting(country,start_date,end_date,category,recipient=None):
    filename = file_name(country=country,start_date=start,end_date=end,category=category,recipient=recipient)
    app.app_context().push()
    db.session.remove()
    db.engine.dispose()
    reporting = (db.session.query(Document.doc_id,
                            Document.date,
                            Document.title,
                            Document.distilled_text,
                            Document.salience_justification,
                            Category.category,
                            TokenizedDocuments.tokens)
                            .join(Category,Document.doc_id==Category.doc_id)
                            .join(InitiatingCountry,Document.doc_id==InitiatingCountry.doc_id)
                            .join(RecipientCountry,Document.doc_id==RecipientCountry.doc_id)
                            .join(TokenizedDocuments,Document.doc_id==TokenizedDocuments.doc_id)
                            .filter(Document.date.between(start,end),
                            InitiatingCountry.initiating_country==country,
                            RecipientCountry.recipient_country.in_(cfg.recipients),
                            Category.category==category).all())
    df_reporting = pd.DataFrame([x for x in reporting])
    # If you know the date column names, e.g., "date"
    df_reporting["date"] = df_reporting["date"].astype(str)
    df_reporting.drop_duplicates(inplace=True)
    df_reporting.reset_index(inplace=True,drop=True)
    df_reporting.to_excel(f'./gai_summary/data/{filename}.xlsx')
    with open(os.path.join(cfg.gai_json,f'{filename}_report_text.json'),'w') as f:
        json.dump(df_reporting.to_dict(orient='records'),f,indent=4)

    str_reporting = ''
    str_reporting += f'TOTAL {country}-{category} REPORTS: {len(reporting)}\n'
    if reporting:
        if len(df_reporting) > 350:
            sample_ = random.sample(reporting, k=275)
            for r in sample_:
                str_reporting += f'{r}\n'
        else:
            for r in reporting:
                str_reporting += f'{r}\n'
    return reporting,str_reporting    

def summary_activities(country,start_date,end_date,category,recipient=None):
    filename = file_name(country=country,start_date=start,end_date=end,category=category,recipient=recipient)
    app.app_context().push()
    db.session.remove()
    db.engine.dispose()
    activities = db.session.query(SoftPowerActivity).filter(SoftPowerActivity.date.between(start,end),
                                                    SoftPowerActivity.initiating_country==country,
                                                    SoftPowerActivity.recipient_country.in_(cfg.recipients),
                                                    SoftPowerActivity.category.like(f"{category}")
                                                    ).all()
    df_ = pd.DataFrame([x.to_dict() for x in activities])
    # If you know the date column names, e.g., "date"
    df_["date"] = df_["date"].astype(str)
    with open(os.path.join(cfg.gai_json,f'{filename}_activities.json'),'w') as f:
        json.dump(df_.to_dict(orient='records'),f,indent=4)
    str_activities = ''
    str_activities += f'ACTIVITY REFERENCES: {len(activities)}\n'
    if activities:
        for a in activities:
            str_activities += f'{a.to_dict()}\n'
    return activities,str_activities

def summary_dailies(country,start_date,end_date,category,recipient=None):
    filename = file_name(country=country,start_date=start,end_date=end,category=category,recipient=recipient)
    app.app_context().push()
    db.session.remove()
    db.engine.dispose()
    dailies = (db.session.query(DailySummary)
                            .join(DailyCategory,DailySummary.id==DailyCategory.daily_id)
                            .filter(DailySummary.date.between(start,end),
                                                        DailySummary.initiating_country==country,
                                                        DailyCategory.category==category,
                                                        DailyRecipient.recipient_country.in_(cfg.recipients)
                                                        ).all()
    )
    df_ = pd.DataFrame([x.to_dict() for x in dailies])
    # If you know the date column names, e.g., "date"
    df_["date"] = df_["date"].astype(str)
    with open(os.path.join(cfg.gai_json,f'{filename}_dailies.json'),'w') as f:
        json.dump(df_.to_dict(orient='records'),f,indent=4)
    str_daily = ''
    if dailies:
        for d in dailies:
            str_daily += f'{d.to_dict()}\n'
    return dailies,str_daily
    
def summary_metrics(country,start_date,end_date,category,recipient=None,dataframe=False):
    app.app_context().push()
    db.session.remove()
    db.engine.dispose()
    metrics = (db.session.query(InitiatingCountry.initiating_country,func.count(InitiatingCountry.doc_id)
                .label("num_docs"))
                .group_by(InitiatingCountry.initiating_country)
                .order_by(func.count(InitiatingCountry.doc_id)
                .desc())
                .join(Document,InitiatingCountry.doc_id == Document.doc_id)
                .join(RecipientCountry,InitiatingCountry.doc_id == RecipientCountry.doc_id)
                .join(Category,Category.doc_id==RecipientCountry.doc_id)
                .filter(InitiatingCountry.initiating_country.in_(cfg.influencers),
                        Document.date.between(start,end),
                        Category.category==category,
                        RecipientCountry.recipient_country.in_(cfg.recipients)).all())
    df = pd.DataFrame(metrics, columns=["initiating_country", "num_docs"])
    df["initiating_country"] = df["initiating_country"].fillna("N/A")
    if dataframe:
        df.to_excel(os.path.join(cfg.gai_data,f'{filename}_metrics.xlsx'),index=False)

    str_metrics = df.to_markdown(index=False)  

    metrics = (db.session.query(RecipientCountry.recipient_country,func.count(RecipientCountry.doc_id)
            .label("num_docs"))
            .group_by(RecipientCountry.recipient_country)
            .order_by(func.count(RecipientCountry.doc_id)
            .desc())
            .join(Document,RecipientCountry.doc_id == Document.doc_id)
            .join(InitiatingCountry,RecipientCountry.doc_id==InitiatingCountry.doc_id)
            .join(Category,Category.doc_id==RecipientCountry.doc_id)
            .filter(RecipientCountry.recipient_country.in_(cfg.recipients),
                    InitiatingCountry.initiating_country==country,
                    Category.category==category,
                    Document.date.between(start,end)).all())
    df = pd.DataFrame(metrics, columns=["recipient_country", "num_docs"])
    df["recipient_country"] = df["recipient_country"].fillna("N/A")
    if dataframe:
        df.to_excel(os.path.join(cfg.gai_data,f'{filename}_rec_metrics.xlsx'),index=False)

    str_recipient_metrics = df.to_markdown(index=False)

    metrics = (db.session.query(Category.category,func.count(Category.doc_id)
            .label("num_docs"))
            .group_by(Category.category)
            .order_by(func.count(Category.doc_id)
            .desc())
            .join(Document,Category.doc_id == Document.doc_id)
            .join(InitiatingCountry,Category.doc_id==InitiatingCountry.doc_id)
            .join(RecipientCountry,Category.doc_id==RecipientCountry.doc_id)
            .filter(RecipientCountry.recipient_country.in_(cfg.recipients),
                    InitiatingCountry.initiating_country==country,
                    Document.date.between(start,end)).all())
    df = pd.DataFrame(metrics, columns=["category", "num_docs"])
    df["category"] = df["category"].fillna("N/A")
    if dataframe:
        df.to_excel(os.path.join(cfg.gai_data,f'{filename}_cat_metrics.xlsx'),index=False)
    str_category_metrics = df.to_markdown(index=False)

    metrics = (db.session.query(Subcategory.subcategory,func.count(Subcategory.doc_id)
            .label("num_docs"))
            .group_by(Subcategory.subcategory)
            .order_by(func.count(Subcategory.doc_id)
            .desc())
            .join(Document,Subcategory.doc_id == Document.doc_id)
            .join(InitiatingCountry,Subcategory.doc_id==InitiatingCountry.doc_id)
            .join(RecipientCountry,Subcategory.doc_id==RecipientCountry.doc_id)
            .join(Category,Category.doc_id==Document.doc_id)
            .filter(RecipientCountry.recipient_country.in_(cfg.recipients),
                    InitiatingCountry.initiating_country==country,
                    Category.category==category,
                    Subcategory.subcategory.in_(cfg.subcategories),
                    Document.date.between(start,end)).all())
    df = pd.DataFrame(metrics, columns=["subcategory", "num_docs"])
    df["subcategory"] = df["subcategory"].fillna("N/A")
    str_subcategory_metrics = df.to_markdown(index=False)
    if dataframe:
            df.to_excel(os.path.join(cfg.gai_data,f'{filename}_subcat_metrics.xlsx'),index=False)

    metrics = (db.session.query(
            func.date_trunc('week', Document.date).label("week_start"),
            func.count(func.distinct(Document.doc_id)).label("num_docs"))
        .join(RecipientCountry, RecipientCountry.doc_id == Document.doc_id)
        .join(InitiatingCountry, InitiatingCountry.doc_id == Document.doc_id)
        .join(Category,Category.doc_id==Document.doc_id)
        .filter(
            InitiatingCountry.initiating_country==country,
            Document.date.between(start, end),
            Category.category==category,
            RecipientCountry.recipient_country.in_(cfg.recipients)
        )
        .group_by("week_start")
        .order_by("week_start")
        .all()
    )
    df = pd.DataFrame(metrics, columns=["week_start", "num_docs"])
    df["week_start"] = df["week_start"].fillna("N/A")
    if dataframe:
            df.to_excel(os.path.join(cfg.gai_data,f'{filename}_WEEKLY_metrics.xlsx'),index=False)
    str_weekly_metrics = df.to_markdown(index=False)

    metrics = (db.session.query(
        func.date_trunc('week', Document.date).label("week_start"),
        func.count(func.distinct(Document.doc_id)).label("num_docs")
    )
    .join(RecipientCountry, RecipientCountry.doc_id == Document.doc_id)
    .join(InitiatingCountry, InitiatingCountry.doc_id == Document.doc_id)
    .join(Category,Category.doc_id==Document.doc_id)
    .filter(
        InitiatingCountry.initiating_country==country,
        Document.date > "2024-08-01",
        Document.date < end,
        Category.category==category,
        RecipientCountry.recipient_country.in_(cfg.recipients)
        )
        .group_by("week_start")
        .order_by("week_start")
        .all()
        )
    df = pd.DataFrame(metrics, columns=["week_start", "num_docs"])
    df["week_start"] = df["week_start"].fillna("N/A")
    if dataframe:
            df.to_excel(os.path.join(cfg.gai_data,f'{filename}_weeklyall_metrics.xlsx'),index=False)
    str_weekly_metrics_all = df.to_markdown(index=False)

    metrics = (db.session.query(
        func.date_trunc('month', Document.date).label("month_start"),
        func.count(func.distinct(Document.doc_id)).label("num_docs")
        )
        .join(RecipientCountry, RecipientCountry.doc_id == Document.doc_id)
        .join(InitiatingCountry, InitiatingCountry.doc_id == Document.doc_id)
        .join(Category,Category.doc_id==Document.doc_id)
        .filter(
            InitiatingCountry.initiating_country==country,
            Document.date > "2024-08-01",
            Document.date < end,
            Category.category==category,
            RecipientCountry.recipient_country.in_(cfg.recipients))
        .group_by("month_start")
        .order_by("month_start")
        .all())
    df = pd.DataFrame(metrics, columns=["month_start", "num_docs"])
    df["month_start"] = df["month_start"].fillna("N/A")
    if dataframe:
            df.to_excel(os.path.join(cfg.gai_data,f'{filename}_monthlyall_metrics.xlsx'),index=False)
    # docs_by_subcategory = df.to_markdown(index=False)
    str_monthly_metrics_all = df.to_markdown(index=False)
    metrics = (db.session.query(
        func.date_trunc('day', Document.date).label("day_published"),
        func.count(Document.doc_id).label("num_docs")
        )
        .join(RecipientCountry, RecipientCountry.doc_id == Document.doc_id)
        .join(InitiatingCountry, InitiatingCountry.doc_id == Document.doc_id)
        .filter(
            InitiatingCountry.initiating_country==country,
            Document.date.between(start, end),
            RecipientCountry.recipient_country.in_(cfg.recipients)
        )
        .group_by("day_published")
        .order_by("day_published")
        .all())
    df = pd.DataFrame(metrics, columns=["day", "num_docs"])
    df["day"] = df["day"].fillna(0)
    if dataframe:
            df.to_excel(os.path.join(cfg.gai_data,f'{filename}_daily_metrics.xlsx'),index=False)
    str_daily_metrics = df.to_markdown(index=False)
    event_metrics = (
    db.session.query(
        Event.event_name,
        func.count(EventSources.doc_id).label("num_docs"))
    .join(EventSources,EventSources.event_id==Event.id)
    .join(Document, EventSources.doc_id == Document.doc_id)
    .join(InitiatingCountry, Document.doc_id == InitiatingCountry.doc_id)
    .join(RecipientCountry, Document.doc_id == RecipientCountry.doc_id)
    .join(Category,Document.doc_id == Category.doc_id)
    .filter(InitiatingCountry.initiating_country==country,
        Document.date.between(start, end),
        Category.category==category,
        RecipientCountry.recipient_country.in_(cfg.recipients))
    .group_by(Event.event_name)
    .order_by(func.count(EventSources.doc_id).desc())
    .limit(20)  # This line limits the results to 20 rows
    .all())
    df = pd.DataFrame(event_metrics, columns=[f"event-{start} to {end}", "num_docs"])
    df[f"event-{start} to {end}"] = df[f"event-{start} to {end}"].fillna("N/A")
    if dataframe:
            df.to_excel(os.path.join(cfg.gai_data,f'{filename}_event_metrics.xlsx'),index=False)
    str_event_metrics = df.to_markdown(index=False)
    
    prompt = ''
    prompt += f'HISTORIC MONTHLY ARTICLE COUNTS:\n{str_monthly_metrics_all}\n\n'
    prompt += f'ARTICLE COUNTS BY COUNTRY:\n{str_metrics}\n\n'
    prompt += f'HISTORIC WEEKLY ARTICLE COUNTS:\n{str_weekly_metrics_all}\n\n'
    prompt += f'CURRENT WEEKLY ARTICLE COUNTS:\n{str_weekly_metrics}\n\n' 
    prompt += f'CURRENT MONTH RECIPIENT METRICS:\n{str_recipient_metrics}\n\n'
    prompt += f'CURRENT MONTH CATEGORY METRICS:\n{str_category_metrics}\n\n' 
    prompt += f'CURRENT MONTH SUBCATEGORY METRICS:\n{str_subcategory_metrics}\n\n' 
    prompt += f'CURRENT MONTH EVENT METRICS:\n{str_event_metrics}\n\n'
    prompt += f'CURRENT MONTH DAILY ARTICLE COUNTS BY DAY:n{str_daily_metrics}\n\n' 
  

    return prompt,str_metrics,str_recipient_metrics,str_category_metrics,str_subcategory_metrics,str_weekly_metrics,str_weekly_metrics_all,str_monthly_metrics_all,str_daily_metrics,str_event_metrics

def insert_summary(country, category, start_date, end_date, event_list, recipient=None, update=False):
    if event_list:
        for event in event_list:
            if recipient:
                # Use upsert or do nothing based on `update`
                if update:
                    stmt = insert(RecipientSummary).values(
                        country=country,
                        category=category,
                        start_date=start_date,
                        end_date=end_date,
                        recipient=recipient,
                        key_event=event['key_event'],
                        overview=event['overview'],
                        outcome=event['outcomes']
                    ).on_conflict_do_update(
                        index_elements=['country', 'key_event', 'start_date', 'end_date', 'category', 'recipient'],
                        set_={
                            "overview": event['overview'],
                            "outcome": event['outcomes'],
                            # Add more fields to update if needed
                        }
                    )
                else:
                    stmt = insert(RecipientSummary).values(
                        country=country,
                        category=category,
                        start_date=start_date,
                        end_date=end_date,
                        recipient=recipient,
                        key_event=event['key_event'],
                        overview=event['overview'],
                        outcome=event['outcomes']
                    ).on_conflict_do_nothing(
                        index_elements=['country', 'key_event', 'start_date', 'end_date', 'category', 'recipient']
                    )
            else:
                if update:
                    stmt = insert(CountrySummary).values(
                        country=country,
                        key_event=event['key_event'],
                        start_date=start_date,
                        end_date=end_date,
                        category=category,
                        overview=event['overview'],
                        outcome=event['outcomes']
                    ).on_conflict_do_update(
                        index_elements=['country', 'key_event', 'start_date', 'end_date', 'category'],
                        set_={
                            "overview": event['overview'],
                            "outcome": event['outcomes'],
                        }
                    )
                else:
                    stmt = insert(CountrySummary).values(
                        country=country,
                        key_event=event['key_event'],
                        start_date=start_date,
                        end_date=end_date,
                        category=category,
                        overview=event['overview'],
                        outcome=event['outcomes']
                    ).on_conflict_do_nothing(
                        index_elements=['country', 'key_event', 'start_date', 'end_date', 'category']
                    )

            db.session.execute(stmt)
        db.session.commit()
    else:
        return 



from backend.scripts.prompts import gai_summary
from backend.scripts.utils import gai,fetch_gai_content
from backend.scripts.models import Document
import datetime
app.app_context().push()
db.session.remove()
db.engine.dispose()

for category in cfg.categories:
    print(f'processing {category}...')
    stmt = (db.session.query(CountrySummary)
            .filter(CountrySummary.country==country,
                    CountrySummary.start_date==start,
                    CountrySummary.end_date==end,
                    CountrySummary.category==category))
    summary_exists = db.session.query(stmt.exists()).scalar()
    # if summary_exists:
    #     print(f'Summary for {country} {category} from {start} to {end} already exists in the database.')
    #     continue
    prompt,str_metrics,str_recipient_metrics,str_category_metrics,str_subcategory_metrics,str_weekly_metrics,str_weekly_metrics_all,str_monthly_metrics_all,str_daily_metrics,str_event_metrics = summary_metrics(country=country,start_date=start,end_date=end,category=category)
    reporting,str_reporting = summary_reporting(country=country,start_date=start,end_date=end,category=category)
    try:
        activities,str_activities = summary_activities(country=country,start_date=start,end_date=end,category=category)
    except:
        print('error processing activities')
        str_activities=''
    try:
        dailies,str_daily_metrics = summary_dailies(country=country,start_date=start,end_date=end,category=category)
    except:
        print('error processing dailies')
        str_daily_metrics = ''
    prompt += f'CURRENT MONTH DAILY SUMMARIES:\n{str_daily_metrics}\n\n'
    prompt += f'CURRENT MONTH Highlighted Activities:\n{str_activities}\n\n'
    prompt += f'CURRENT MONTH REPORTING:\n{str_reporting}\n\n'
    sys_prompt = gai_summary.format(country=country,category=category,start_date=start,end_date=end,date_string=today,top_n=8)
    response = gai(sys_prompt=sys_prompt,user_prompt=prompt)
    gai_output = fetch_gai_content(response)
    print('inserting summary')
    insert_summary(country=country,category=category,start_date=start,end_date=end,event_list=gai_output,update=True)



from backend.scripts.prompts import gai_summary
from backend.scripts.utils import gai,fetch_gai_content
from backend.scripts.models import Document
app.app_context().push()
db.session.remove()
db.engine.dispose() 
exclude_sources = ["CRIonline Arabic", 
                    "CGTN Online (Arabic)", 
                    "China in Arab Eyes Online", 
                    "People's Daily Online", 
                    "China Silk Road News Network", 
                    "China Today Online", 
                    "Xinhuanet (Arabic)", 
                    "China Arab TV",
                    "Belt and Road Portal", 
                    "China in Arabic", 
                    "OSAS News China-Arab Media Agency"]

for category in cfg.categories:
    reporting = (db.session.query(Document.doc_id,
                            Document.date,
                            Document.source_name,
                            Document.title,
                            Document.distilled_text,
                            Document.salience_justification,
                            Category.category,
                            TokenizedDocuments.tokens)
                            .join(Category,Document.doc_id==Category.doc_id)
                            .join(InitiatingCountry,Document.doc_id==InitiatingCountry.doc_id)
                            .join(RecipientCountry,Document.doc_id==RecipientCountry.doc_id)
                            .join(TokenizedDocuments,Document.doc_id==TokenizedDocuments.doc_id)
                            .filter(Document.date.between(start,end),
                            Document.source_name.not_in(exclude_sources),
                            InitiatingCountry.initiating_country==country,
                            RecipientCountry.recipient_country.in_(cfg.recipients),
                            Category.category==category).all())
    articles = pd.DataFrame([x for x in reporting])


    summaries = (db.session.query(CountrySummary)
                    .filter(CountrySummary.country==country,
                    CountrySummary.start_date==start,
                    CountrySummary.end_date==end,
                    CountrySummary.category==category))

    events = pd.DataFrame([x.to_dict() for x in summaries])
    events['description'] = [f"{e['key_event']} {e['overview']} {e['outcome']}" for index,e in events.iterrows()]
    events['processed_description'] = events['description'].apply(preprocess)
    vectorizer = TfidfVectorizer()
    article_vectors = vectorizer.fit_transform(articles['tokens'])
    # Example usage


    for index, event in events.iterrows():
        event_description = event['processed_description']
        similarity_scores = find_related_articles(event_description, article_vectors, vectorizer, articles)
        articles['similarity_score'] = similarity_scores
        # Filter articles based on a similarity threshold
        threshold = 0.15  # Adjust threshold as needed
        relevant_articles = articles[articles['similarity_score'] > threshold]
        
        overview_output = fetch_gai_sources(source_summary(records=relevant_articles.to_dict(orient='records'),
                                                        summary=event.to_dict(),
                                                        summ_type='overview'))
        if not overview_output:
            continue
        print(f"{event['key_event']} overview:{overview_output}")
        outcome_output = fetch_gai_sources(source_summary(records=relevant_articles.to_dict(orient='records'),
                                        summary=event.to_dict(),
                                        summ_type='outcome'))
        if not outcome_output:
            continue
        print(f"{event['key_event']} outcome:{outcome_output}")
        for source in overview_output:
            stmt = insert(SummarySources).values(
                        summary_id = event['id'], 
                        summary_type ='overview',  
                        doc_id = source
                    ).on_conflict_do_update(
                        index_elements=['summary_id','doc_id'],
                        set_={
                            "summary_id": event['id'],
                            "summary_type": 'overview',
                            "doc_id": source,
                            # Add more fields to update if needed
                        })
            db.session.execute(stmt)
        db.session.commit()
        for source in outcome_output:
            stmt = insert(SummarySources).values(
                        summary_id = event['id'],
                        summary_type = 'outcome',
                        doc_id = source
                    ).on_conflict_do_update(
                        index_elements=['summary_id','doc_id'],
                        set_={
                            "summary_id": event['id'],
                            "summary_type": 'outcome',
                            "doc_id": source,
                            # Add more fields to update if needed
                        })
            db.session.execute(stmt)
        db.session.commit()
        print(f'processed {event["id"]}: {event["key_event"]}')

from backend.scripts.utils import build_hyperlink
import re
with app.app_context():
    for category in cfg.categories:
        summaries = (db.session.query(CountrySummary)
                        .filter(CountrySummary.country==country,
                        CountrySummary.start_date==start,
                        CountrySummary.end_date==end,
                        CountrySummary.category==category))
        summaries = [x.to_dict() for x in summaries]
        for s in summaries:
                sources =(db.session.query(SummarySources.doc_id,SummarySources.summary_type)
                        .join(CountrySummary,SummarySources.summary_id==CountrySummary.id)
                        .filter(CountrySummary.id==s['id'])).all()
                outcome_sources = [x[0] for x in sources if x[1]=='outcome']
                outcome_sources = [clean_string(x) for x in outcome_sources]
                overview_sources = [x[0] for x in sources if x[1]=='overview']
                overview_sources = [clean_string(x) for x in overview_sources]
                s['summary_citations'] = overview_sources
                s['outcome_citations'] = outcome_sources
                s['summary_hyperlink'] = build_hyperlink(overview_sources)
                s['outcome_hyperlink'] = build_hyperlink(outcome_sources)
        
        summary_df = pd.DataFrame(summaries)
        summary_df = summary_df[summary_df['summary_citations'].apply(len)>0]
        summary_df = summary_df[summary_df['outcome_citations'].apply(len)>0]
        summaries = summary_df.to_dict(orient='records')
        summary_dict = {}
        for s in summaries:
                summary_dict[s['id']] = {}
                for k,v in s.items():
                        if k != 'id':
                                summary_dict[s['id']][k] = v
                                
        filename = file_name(country=country,category=category,start_date=start,end_date=end,recipient=recipient)
        directory = cfg.gai_json
        with open(os.path.join(directory,f'{filename}summaries.json'),'w') as f:
                json.dump(summary_dict,f,indent=4)

def consolidate_summaries(summary_dict,country,start_date,end_date):
    directory = cfg.gai_json
    filename = file_name(country=country,start_date=start_date,end_date=end_date,category=None,recipient=None)
    summaries = {}
    for category in cfg.categories:
        file = f'{filename}_{category}summaries.json'
        loaded_summaries = json.load(open(os.path.join(directory,file)))
        for k,v in loaded_summaries.items():
            summaries[k] = v
    with open(os.path.join(directory,f'{filename}summaries.json'),'w') as f:
        json.dump(summaries,f,indent=4)

consolidate_summaries(summary_dict=summary_dict,country=country,start_date=start,end_date=end)


from backend.scripts.prompts import consolidation_prompt
from backend.scripts.utils import gai,fetch_gai_content
def event_summaries_outcomes(country,start_date,end_date,recipient=None):
    filename = file_name(country,start_date,end_date,category=None,recipient=recipient)
    directory = cfg.gai_json
    summaries = json.load(open(os.path.join(directory,f"{filename}summaries.json")))
    ids = list(summaries.keys())
    events = {}
    for category in cfg.categories:
        events[category] = []
        for id_ in ids:
            x = summaries[id_]
            if x['category'] == category:
                event = {}
                event['id'] = id_
                event['event_name'] = x['key_event']
                event['content'] = x['overview'] + x['outcome']
                events[category].append(event)
    return events

def process_output(output):
    if isinstance(output,dict):
        return output
    if isinstance(output,list):
        return output[0]

def deduplicate_events(country,start_date,end_date,category=None,recipient=None):
    filename = file_name(country,start_date,end_date,category=category,recipient=recipient)
    directory = cfg.gai_json
    summaries = json.load(open(os.path.join(directory,f"{filename}summaries.json")))
    events = event_summaries_outcomes(country,start_date,end_date,recipient=recipient)
    consolidation_response = gai(consolidation_prompt,str(events))
    c_response = fetch_gai_content(consolidation_response)
    result = process_output(c_response)
    with open(os.path.join(directory,f"{filename}deduplication.json"),'w') as f:
        json.dump(result,f,indent=4)
    mapping = {}
    ids = []
    for k,v in result.items():
        mapping[k] = []
        for id_ in v:
            mapping[k].append(summaries[id_]['key_event'])
            ids.append(id_)
    mapping['unlisted'] = [summaries[id_]['key_event'] for id_ in list(summaries.keys()) if id_ not in ids]
    with open(os.path.join(directory,f"{filename}mapping.json"),'w') as f:
        json.dump(mapping,f,indent=4)
deduplicate_events(country=country,start_date=start,end_date=end)


from docx import Document
from datetime import datetime
from docx.shared import Pt
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
import json

def format_date_range(start_date_str, end_date_str):
    start = datetime.strptime(start_date_str, "%Y-%m-%d")
    end = datetime.strptime(end_date_str, "%Y-%m-%d")

    start_formatted = start.strftime("%B %Y")
    end_formatted = end.strftime("%B %Y")

    if start_formatted == end_formatted:
        return start_formatted
    else:
        return f"{start_formatted} to {end_formatted}"

def sample_summaries(summaries,perc=.65):
    sorted_summaries = sorted(summaries)
    sample_size = int(len([key for key in summaries.keys()])*perc)
    random.seed(42)  
    sampled_docs = random.sample(sorted_summaries, sample_size)
    sample = {k:summaries[k] for k in sampled_docs}
    return sample

def format_full_date_range(start_date_str, end_date_str):
    start = datetime.strptime(start_date_str, "%Y-%m-%d")
    end = datetime.strptime(end_date_str, "%Y-%m-%d")

    start_formatted = start.strftime("%-d %B %Y")  # e.g., 1 February 2025
    end_formatted = end.strftime("%-d %B %Y")      # e.g., 28 February 2025

    return f"{start_formatted} to {end_formatted}"

def replace_placeholders_in_para(para, mapping):
    # if the full placeholder appears anywhere in this paragraph...
    for key, val in mapping.items():
        if key in para.text:
            # grab the full text, do the replace, then re‑build with one run
            new_text = para.text.replace(key, val)
            # clear out existing runs
            for run in para.runs:
                run.text = ""
            # put in one run with the replaced text
            para.add_run(new_text)
          

def insert_paragraph_after(existing_para, style_name=None):
    """
    Insert a new <w:p> immediately after `existing_para`, optionally set its style.
    Returns the new Paragraph.
    """
    new_p_elm = OxmlElement('w:p')
    existing_para._p.addnext(new_p_elm)
    new_para = Paragraph(new_p_elm, existing_para._parent)

    # only apply if that style actually exists in the document
    if style_name and style_name in {s.name for s in existing_para.part.styles}:
        new_para.style = style_name

    return new_para

def replace_in_block(block, mapping):
    # block can be a Document, Header, Footer, Cell, etc.
    for para in block.paragraphs:
        replace_placeholders_in_para(para, mapping)
    for table in getattr(block, 'tables', []):
        for row in table.rows:
            for cell in row.cells:
                replace_in_block(cell, mapping)

def build_title(summaries,start_date,end_date):
    sys_prompt = f''' Review the following summaries and create an appropriate descriptive title that captures the relevant content, is no more 10 words, and does not qualify or use subjective language. Ensure the title references the date range {format_full_date_range(start_date_str=start_date,end_date_str=end_date)}'''
    user_prompt = str(summaries)
    try:
        response = gai(sys_prompt=sys_prompt,user_prompt=user_prompt)
    
    except:
        print('prompt error: sampling summaries...')
        sample = sample_summaries(summaries)
        user_prompt = str(sample)
        response = gai(sys_prompt=sys_prompt,user_prompt=user_prompt)
    
    title= response['choices'][0]['message']['content'] 
    return title


from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_hyperlink(paragraph, url, text):
    """Insert a clickable external hyperlink into `paragraph`."""
    part = paragraph.part
    # create relationship
    r_id = part.relate_to(url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True)
    # build w:hyperlink element
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # build a run inside it
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    # add blue + underline
    c = OxmlElement('w:color');  c.set(qn('w:val'), "0000FF")
    u = OxmlElement('w:u');      u.set(qn('w:val'), "single")
    rPr.append(c); rPr.append(u)
    new_run.append(rPr)

    # text
    t = OxmlElement('w:t'); t.text = text
    new_run.append(t)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return paragraph


def add_image_hyperlink(paragraph, image_path, target_url, width=None, height=None):
    """
    In `paragraph`, insert a picture (via add_picture) and immediately
    wrap it in a clickable external hyperlink to `target_url`.
    Returns the paragraph for chaining.
    """
    part = paragraph.part

    # 1) Insert the picture as a new run
    pic_run = paragraph.add_run()
    # use Inches or EMU conversion as you like
    pic_run.add_picture(image_path, width=Inches(width) if width else None,
                                   height=Inches(height) if height else None)

    # 2) Build the hyperlink rel & element
    rId_link = part.relate_to(
        target_url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True
    )
    hlink = OxmlElement('w:hyperlink')
    hlink.set(qn('r:id'), rId_link)

    # 3) Pluck out that last <w:r> (the picture run) and reparent it under <w:hyperlink>
    run_elm = paragraph._p[-1]                 # last child is the <w:r> we just added
    paragraph._p.remove(run_elm)               # remove it
    hlink.append(run_elm)                      # put it inside the hyperlink
    paragraph._p.append(hlink)                 # re-append hyperlink to the paragraph

    return paragraph

    def reviewer_version(country,start_date,end_date):
    # Load the template document
    doc = Document('./gai_summary/GAI_Summary_Template.docx')
    filename = file_name(country=country,start_date=start,end_date=end,category=None,recipient=None)
    report_ids = json.load(open(os.path.join(cfg.gai_json,f'{filename}deduplication.json')))
    summaries = json.load(open(os.path.join(cfg.gai_json,f'{filename}summaries.json')))
    # sources = json.load(open(os.path.join(cfg.gai_json,f'{filename}sources.json')))
    title = build_title(summaries,start_date=start_date,end_date=end_date)
    # pre‑cache available style names
    available_styles = {s.name for s in doc.styles}

    placeholders = {
        'Economic': '{{economic_event_section}}',
        'Diplomacy': '{{diplomatic_event_section}}',
        'Social':   '{{social_event_section}}',
        'Military':'{{military_event_section}}'
    }
    global_placeholders = {
        '{{country}}': country,
        '{{date}}': format_full_date_range(start_date_str=start_date,end_date_str=end_date),
        '{{summary_title}}': title
    }


    for para in doc.paragraphs:
        replace_placeholders_in_para(para, global_placeholders)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_placeholders_in_para(para, global_placeholders)
    for section in doc.sections:
        for para in section.header.paragraphs:
            replace_placeholders_in_para(para, global_placeholders)
        for para in section.footer.paragraphs:
            replace_placeholders_in_para(para, global_placeholders)

    summary_ids = list(summaries.keys())

    for category in cfg.categories:
        tag = placeholders[category]

        # find the one paragraph that contains this placeholder
        for para in doc.paragraphs:
            if tag in para.text:
                # clear the placeholder text
                para.text = ''
                last = para  # anchor point for insertion

                # for each summary in this category
                for sid in summary_ids:
                    x = summaries[sid]
                    if x['category'] != category:
                        continue
                    if sid not in report_ids[category]:
                        continue
                    # 1) Event heading
                    ev = insert_paragraph_after(last,
                                            style_name='Heading 3' if 'Heading 3' in available_styles else None)
                    run = ev.add_run(x['key_event'])
                    run.bold = True

                    last = ev
                    # 2) Overview
                    ov = insert_paragraph_after(last, style_name='Normal')
                    # first run: the “Outcomes:” label, in bold
                    label_run = ov.add_run("Overview: ")
                    label_run.bold = True
                    # second run: the actual outcome text, in normal weight
                    text_run = ov.add_run(x['overview'])
                    ov.add_run(" ")
                    # ✅ wrap *this same paragraph* in the hyperlink
                    if x.get('summary_hyperlink'):
                        last = add_image_hyperlink(
                            ov,    # ← use the Overview paragraph itself
                            image_path=os.path.join('./gai_summary', 'img/atom.png'),
                            target_url=x['summary_hyperlink'],
                            width=0.2
                        )
                    last = ov
                    # 3) Summary citations
                    for cite_text in x['summary_citations']:
                        c = insert_paragraph_after(last, style_name='Normal')
                        fmt = c.paragraph_format
                        fmt.space_before = Pt(0)
                        fmt.space_after  = Pt(0)
                        fmt.line_spacing = 1  # single‑line
                        run = c.add_run(cite_text)
                        run.font.size = Pt(8)
                        last = c

                    # 4) Outcomes
                    out = insert_paragraph_after(last, style_name='Normal')
                    # first run: the “Outcomes:” label, in bold
                    label_run = out.add_run("Outcomes: ")
                    label_run.bold = True

                    # second run: the actual outcome text, in normal weight
                    text_run = out.add_run(x['outcome'])
                    out.add_run(" ")
                    # ✅ wrap *this same paragraph* in the hyperlink
                    if x.get('summary_hyperlink'):
                        last = add_image_hyperlink(
                            out,    # ← use the Overview paragraph itself
                            image_path=os.path.join('./gai_summary','img/atom.png'),
                            target_url=x['outcome_hyperlink'],
                            width=0.2
                        )
                    last = out
                    # 5) Outcome citations
                    for cite_text in x['outcome_citations']:
                        c = insert_paragraph_after(last, style_name='Normal')
                        fmt = c.paragraph_format
                        fmt.space_before = Pt(0)
                        fmt.space_after  = Pt(0)
                        fmt.line_spacing = 1
                        run = c.add_run(cite_text)
                        run.font.size = Pt(8)
                        last = c

                break  # done with this placeholder

    doc.save(f'./gai_summary/{country}_{start_date}_{end_date}_Reviewer.docx')
    
def summary_version(country,start_date,end_date):
    # Load the template document
    doc = Document('./gai_summary/GAI_Summary_Template.docx')

    filename = file_name(country=country,start_date=start,end_date=end,category=None,recipient=None)
    report_ids = json.load(open(os.path.join(cfg.gai_json,f'{filename}deduplication.json')))
    summaries = json.load(open(os.path.join(cfg.gai_json,f'{filename}summaries.json')))
    # sources = json.load(open(os.path.join(cfg.gai_json,f'{filename}sources.json')))
    title = build_title(summaries,start_date=start_date,end_date=end_date)
    
    global_placeholders = {
        '{{country}}': country,
        '{{date}}':    format_full_date_range(start_date_str=start_date,
                                              end_date_str=end_date),
        '{{summary_title}}': title
    }
    replace_in_block(doc, global_placeholders)
    for section in doc.sections:
        replace_in_block(section.header, global_placeholders)
        replace_in_block(section.footer, global_placeholders)

    available_styles = {s.name for s in doc.styles}

    placeholders = {
        'Economic':  '{{economic_event_section}}',
        'Diplomacy': '{{diplomatic_event_section}}',
        'Social':    '{{social_event_section}}',
        'Military':  '{{military_event_section}}'
    }

    # — Collect citations by category→event —
    event_citations = {cat: {} for cat in cfg.categories}
    summary_ids     = list(summaries.keys())

    for category in cfg.categories:
        tag = placeholders[category]
        for para in doc.paragraphs:
            if tag in para.text:
                para.text = ''
                last = para
                for sid in summary_ids:
                    x = summaries[sid]
                    if x['category'] != category:
                        continue
                    if sid not in report_ids[category]:
                        continue
                    event = x['key_event']
                    # render event
                    ev = insert_paragraph_after(
                        last,
                        style_name='Heading 3' if 'Heading 3' in available_styles else None
                    )
                    ev.add_run(event).bold = True
                    last = ev

                    # overview
                    ov = insert_paragraph_after(last, style_name='Normal')
                    ov.add_run("Overview: ").bold = True
                    ov.add_run(x['overview'])
                    last = ov

                    # outcomes
                    out = insert_paragraph_after(last, style_name='Normal')
                    out.add_run("Outcomes: ").bold = True
                    out.add_run(x['outcome'])
                    last = out

                    # collect cites
                    bucket = event_citations[category].setdefault(event, [])
                    bucket.extend(x['summary_citations'])
                    bucket.extend(x['outcome_citations'])

                break

    # — Append End Notes with deduplication —
    # Main heading
    if 'Heading 1' in available_styles:
        hn = doc.add_paragraph(style='Heading 1')
    else:
        hn = doc.add_paragraph()
    hn.add_run("End Notes").bold = True

    for category in cfg.categories:
        # category heading
        if 'Heading 2' in available_styles:
            ch = doc.add_paragraph(style='Heading 2')
        else:
            ch = doc.add_paragraph()
        ch.add_run(category).bold = True

        for event, cites in event_citations[category].items():
            # event sub‑heading
            if 'Heading 3' in available_styles:
                eh = doc.add_paragraph(style='Heading 3')
            else:
                eh = doc.add_paragraph()
            eh.add_run(event).bold = True

            # dedupe while preserving order
            seen = []
            for cite in cites:
                if cite not in seen:
                    seen.append(cite)

            # write each unique citation
            for cite in seen:
                p = doc.add_paragraph(style='Normal')
                fmt = p.paragraph_format
                fmt.space_before = Pt(0)
                fmt.space_after  = Pt(0)
                fmt.line_spacing = 1
                run = p.add_run(cite)
                run.font.size = Pt(8)

    # — Save —
    output_path = f'./gai_summary/{country}_{start_date}_{end_date}_Summary.docx'
    doc.save(output_path)

reviewer_version(country,start,end)
summary_version(country,start,end)