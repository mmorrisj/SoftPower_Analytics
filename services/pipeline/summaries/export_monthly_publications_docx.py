#!/usr/bin/env python3
"""
Export monthly summary publications as DOCX files with accompanying CSV source files.

This script generates publication-ready DOCX documents from monthly event summaries
and creates CSV files listing all source documents for validation purposes.

Usage:
    python export_monthly_publications_docx.py --country China --month 2024-08
    python export_monthly_publications_docx.py --influencers --month 2024-08
    python export_monthly_publications_docx.py --country "United States" --start-date 2024-08-01 --end-date 2024-09-30
"""

import argparse
import json
from pathlib import Path
from datetime import datetime, date
from typing import List, Dict, Set
from collections import defaultdict

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd
from sqlalchemy import text

from shared.database.database import get_session
from shared.models.models import EventSummary, EventSourceLink, PeriodType, EventStatus, Document as DocModel
import yaml


def get_monthly_summaries(session, country: str, start_date: date, end_date: date) -> List[Dict]:
    """Get all monthly event summaries for a country and time period."""
    summaries = session.query(EventSummary).filter(
        EventSummary.initiating_country == country,
        EventSummary.period_type == PeriodType.MONTHLY,
        EventSummary.period_start >= start_date,
        EventSummary.period_end <= end_date,
        EventSummary.status == EventStatus.ACTIVE,
        EventSummary.is_deleted == False
    ).order_by(EventSummary.event_name).all()

    results = []
    for summary in summaries:
        # Extract narrative sections from JSONB
        narrative = summary.narrative_summary or {}
        overview = narrative.get('monthly_overview', narrative.get('overview', ''))
        outcome = narrative.get('key_outcomes', narrative.get('outcome', ''))
        strategic_significance = narrative.get('strategic_significance', '')

        # Get source document IDs
        source_links = session.query(EventSourceLink).filter(
            EventSourceLink.event_summary_id == summary.id
        ).all()
        doc_ids = [link.doc_id for link in source_links]

        # Extract metadata
        count_by_category = summary.count_by_category or {}
        count_by_recipient = summary.count_by_recipient or {}
        count_by_source = summary.count_by_source or {}

        results.append({
            'event_name': summary.event_name,
            'overview': overview,
            'outcome': outcome,
            'strategic_significance': strategic_significance,
            'source_doc_ids': doc_ids,
            'period_start': summary.period_start.isoformat() if summary.period_start else None,
            'period_end': summary.period_end.isoformat() if summary.period_end else None,
            'metadata': {
                'categories': count_by_category,
                'recipients': count_by_recipient,
                'top_sources': count_by_source
            }
        })

    # Sort events by number of source documents (descending)
    results.sort(key=lambda x: len(x['source_doc_ids']), reverse=True)

    return results


def get_source_documents(session, doc_ids: List[str]) -> pd.DataFrame:
    """Get source document details for a list of doc_ids."""
    if not doc_ids:
        return pd.DataFrame(columns=['doc_id', 'title', 'distilled_text', 'date', 'source'])

    # Query documents
    docs = session.query(DocModel).filter(
        DocModel.doc_id.in_(doc_ids)
    ).all()

    # Build DataFrame
    data = []
    for doc in docs:
        data.append({
            'doc_id': doc.doc_id,
            'title': doc.title or '',
            'distilled_text': doc.distilled_text or '',
            'date': doc.date.isoformat() if doc.date else '',
            'source': doc.source_name or ''
        })

    df = pd.DataFrame(data)
    return df


def create_docx_publication(country: str, start_date: date, end_date: date, events: List[Dict], output_path: Path):
    """Create a DOCX publication from monthly event summaries."""
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Title
    title = doc.add_heading(f'{country} Monthly Summary', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Period subtitle
    period_text = f'{start_date.strftime("%B %Y")}'
    if start_date.month != end_date.month or start_date.year != end_date.year:
        period_text = f'{start_date.strftime("%B %Y")} - {end_date.strftime("%B %Y")}'

    subtitle = doc.add_paragraph(period_text)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.runs[0]
    subtitle_format.font.size = Pt(14)
    subtitle_format.font.color.rgb = RGBColor(89, 89, 89)

    # Generated date
    gen_date = doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
    gen_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    gen_date_format = gen_date.runs[0]
    gen_date_format.font.size = Pt(10)
    gen_date_format.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph()  # Spacing

    # Summary statistics
    doc.add_heading('Summary', 1)
    doc.add_paragraph(f'Total Events: {len(events)}')

    # Count events by category (from metadata)
    category_counts = defaultdict(int)
    for event in events:
        categories = event.get('metadata', {}).get('categories', {})
        if not categories:
            category_counts['Uncategorized'] += 1
        else:
            for cat in categories.keys():
                category_counts[cat] += 1

    doc.add_paragraph('Events by Category:')
    for cat, count in sorted(category_counts.items()):
        doc.add_paragraph(f'  • {cat}: {count}', style='List Bullet 2')

    doc.add_page_break()

    # Events
    doc.add_heading('Events', 1)

    for i, event in enumerate(events, 1):
        # Event name
        event_heading = doc.add_heading(f'{i}. {event["event_name"]}', 2)

        # Overview section
        if event['overview']:
            doc.add_heading('Overview', 3)
            para = doc.add_paragraph(event['overview'])

            # Add source doc IDs as footnote
            if event['source_doc_ids']:
                source_para = doc.add_paragraph()
                source_run = source_para.add_run(f'[Sources: {len(event["source_doc_ids"])} documents]')
                source_run.font.size = Pt(9)
                source_run.font.color.rgb = RGBColor(128, 128, 128)
                source_run.italic = True

        # Outcome section
        if event['outcome']:
            doc.add_heading('Key Outcomes', 3)
            para = doc.add_paragraph(event['outcome'])

            if event['source_doc_ids']:
                source_para = doc.add_paragraph()
                source_run = source_para.add_run(f'[Sources: {len(event["source_doc_ids"])} documents]')
                source_run.font.size = Pt(9)
                source_run.font.color.rgb = RGBColor(128, 128, 128)
                source_run.italic = True

        # Strategic Significance section
        if event['strategic_significance']:
            doc.add_heading('Strategic Significance', 3)
            para = doc.add_paragraph(event['strategic_significance'])

            if event['source_doc_ids']:
                source_para = doc.add_paragraph()
                source_run = source_para.add_run(f'[Sources: {len(event["source_doc_ids"])} documents]')
                source_run.font.size = Pt(9)
                source_run.font.color.rgb = RGBColor(128, 128, 128)
                source_run.italic = True

        # Metadata
        metadata = event.get('metadata', {})
        if metadata.get('recipients') or metadata.get('categories'):
            doc.add_heading('Metadata', 3)

            if metadata.get('recipients'):
                doc.add_paragraph('Recipients:', style='List Bullet')
                for recipient, count in sorted(metadata['recipients'].items()):
                    doc.add_paragraph(f'{recipient}: {count} documents', style='List Bullet 2')

            if metadata.get('categories'):
                doc.add_paragraph('Categories:', style='List Bullet')
                for category, count in sorted(metadata['categories'].items()):
                    doc.add_paragraph(f'{category}: {count} documents', style='List Bullet 2')

        doc.add_paragraph()  # Spacing between events

    # Save document
    doc.save(output_path)
    print(f"[SAVED] {output_path}")


def create_source_csv(country: str, start_date: date, end_date: date, events: List[Dict], session, output_path: Path):
    """Create a CSV file with all source documents referenced in the publication."""
    # Collect all unique doc_ids
    all_doc_ids = set()
    event_doc_mapping = {}  # Track which events use which docs

    for event in events:
        event_name = event['event_name']
        doc_ids = event['source_doc_ids']
        all_doc_ids.update(doc_ids)

        for doc_id in doc_ids:
            if doc_id not in event_doc_mapping:
                event_doc_mapping[doc_id] = []
            event_doc_mapping[doc_id].append(event_name)

    # Get source documents
    if not all_doc_ids:
        print("[WARNING] No source documents found")
        return

    df = get_source_documents(session, list(all_doc_ids))

    # Add event names column
    df['events'] = df['doc_id'].apply(lambda x: '; '.join(event_doc_mapping.get(x, [])))

    # Reorder columns
    df = df[['doc_id', 'title', 'date', 'source', 'events', 'distilled_text']]

    # Sort by date descending
    df['date'] = pd.to_datetime(df['date'], errors='coerce')
    df = df.sort_values('date', ascending=False)
    df['date'] = df['date'].dt.strftime('%Y-%m-%d')

    # Save CSV
    df.to_csv(output_path, index=False, encoding='utf-8-sig')
    print(f"[SAVED] {output_path} ({len(df)} documents)")


def export_publications(countries: List[str], start_date: date, end_date: date, output_dir: Path):
    """Export DOCX publications and CSV source files for specified countries."""
    output_dir.mkdir(parents=True, exist_ok=True)

    with get_session() as session:
        for country in countries:
            print(f"\nProcessing {country}...")

            # Get monthly summaries
            events = get_monthly_summaries(session, country, start_date, end_date)

            if not events:
                print(f"  [SKIP] No monthly summaries found for {country}")
                continue

            print(f"  Found {len(events)} monthly event summaries")

            # Generate filename base
            filename_base = f"{country.replace(' ', '_')}_{start_date.isoformat()}_{end_date.isoformat()}"

            # Create DOCX publication
            docx_path = output_dir / f"{filename_base}_monthly_summary.docx"
            create_docx_publication(country, start_date, end_date, events, docx_path)

            # Create CSV source file
            csv_path = output_dir / f"{filename_base}_sources.csv"
            create_source_csv(country, start_date, end_date, events, session, csv_path)

            print(f"  [COMPLETE] {country}")


def main():
    parser = argparse.ArgumentParser(
        description='Export monthly summary publications as DOCX with CSV source files'
    )
    parser.add_argument('--country', type=str, help='Single country to process')
    parser.add_argument('--influencers', action='store_true', help='Process all influencer countries from config')
    parser.add_argument('--month', type=str, help='Month in YYYY-MM format (uses full month)')
    parser.add_argument('--start-date', type=str, help='Start date (YYYY-MM-DD)')
    parser.add_argument('--end-date', type=str, help='End date (YYYY-MM-DD)')
    parser.add_argument('--output-dir', type=str, default='_data/publications',
                       help='Output directory for DOCX and CSV files')

    args = parser.parse_args()

    # Determine countries
    countries = []
    if args.influencers:
        with open("shared/config/config.yaml") as f:
            countries = yaml.safe_load(f).get("influencers", [])
    elif args.country:
        countries = [args.country]
    else:
        print("Error: Must specify --country or --influencers")
        return

    # Determine date range
    if args.month:
        # Parse month (YYYY-MM)
        year, month = map(int, args.month.split('-'))
        start_date = date(year, month, 1)
        # Get last day of month
        if month == 12:
            end_date = date(year + 1, 1, 1)
        else:
            end_date = date(year, month + 1, 1)
        from datetime import timedelta
        end_date = end_date - timedelta(days=1)
    elif args.start_date and args.end_date:
        start_date = date.fromisoformat(args.start_date)
        end_date = date.fromisoformat(args.end_date)
    else:
        print("Error: Must specify --month OR --start-date and --end-date")
        return

    output_dir = Path(args.output_dir)

    print("=" * 80)
    print("EXPORTING MONTHLY PUBLICATIONS (DOCX + CSV)")
    print("=" * 80)
    print(f"Countries: {', '.join(countries)}")
    print(f"Period: {start_date} to {end_date}")
    print(f"Output: {output_dir}")
    print("=" * 80)
    print()

    export_publications(countries, start_date, end_date, output_dir)

    print()
    print("=" * 80)
    print("EXPORT COMPLETE")
    print("=" * 80)


if __name__ == "__main__":
    main()
