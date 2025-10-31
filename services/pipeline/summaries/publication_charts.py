#!/usr/bin/env python3
"""
Chart generation module for DOCX publications.

Creates publication-ready charts and embeds them into Word documents.
"""

import io
from pathlib import Path
from typing import Dict, List
from collections import Counter

import matplotlib
matplotlib.use('Agg')  # Non-interactive backend for server use
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from docx.shared import Inches


def create_category_pie_chart(events_by_category: Dict[str, List[Dict]]) -> io.BytesIO:
    """
    Create a pie chart showing event distribution by category.

    Args:
        events_by_category: Dict with keys 'Economic', 'Diplomatic', 'Social'

    Returns:
        BytesIO object containing the chart image
    """
    # Count events per category
    counts = {cat: len(events) for cat, events in events_by_category.items() if events}

    if not counts:
        return None

    # Create pie chart
    fig, ax = plt.subplots(figsize=(8, 6))
    colors = ['#2E86AB', '#A23B72', '#F18F01']  # Blue, Purple, Orange

    wedges, texts, autotexts = ax.pie(
        counts.values(),
        labels=counts.keys(),
        autopct='%1.1f%%',
        startangle=90,
        colors=colors,
        textprops={'fontsize': 12, 'weight': 'bold'}
    )

    # Make percentage text white
    for autotext in autotexts:
        autotext.set_color('white')
        autotext.set_fontsize(14)

    ax.set_title('Event Distribution by Category', fontsize=16, weight='bold', pad=20)
    plt.tight_layout()

    # Save to BytesIO
    img_stream = io.BytesIO()
    plt.savefig(img_stream, format='png', dpi=300, bbox_inches='tight')
    img_stream.seek(0)
    plt.close()

    return img_stream


def create_top_recipients_chart(events_by_category: Dict[str, List[Dict]], top_n: int = 10) -> io.BytesIO:
    """
    Create a horizontal bar chart showing top recipient countries.

    Args:
        events_by_category: Dict with event data
        top_n: Number of top recipients to show

    Returns:
        BytesIO object containing the chart image
    """
    # Aggregate all recipient counts across all events
    recipient_counts = Counter()

    for category, events in events_by_category.items():
        for event in events:
            # Each event should have source_doc_ids - we'll use that as a proxy
            # In practice, you'd extract from count_by_recipient if available
            recipient_counts[f"Event: {event['event_name'][:30]}..."] += len(event.get('source_doc_ids', []))

    if not recipient_counts:
        return None

    # Get top N
    top_recipients = recipient_counts.most_common(top_n)
    countries = [r[0] for r in top_recipients]
    counts = [r[1] for r in top_recipients]

    # Create horizontal bar chart
    fig, ax = plt.subplots(figsize=(10, max(6, len(countries) * 0.4)))

    bars = ax.barh(countries, counts, color='#2E86AB')

    # Add value labels on bars
    for i, (bar, count) in enumerate(zip(bars, counts)):
        ax.text(count + 0.5, i, str(count), va='center', fontsize=10, weight='bold')

    ax.set_xlabel('Number of Source Documents', fontsize=12, weight='bold')
    ax.set_title(f'Top {len(countries)} Events by Coverage', fontsize=16, weight='bold', pad=20)
    ax.invert_yaxis()  # Highest at top
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)

    plt.tight_layout()

    # Save to BytesIO
    img_stream = io.BytesIO()
    plt.savefig(img_stream, format='png', dpi=300, bbox_inches='tight')
    img_stream.seek(0)
    plt.close()

    return img_stream


def create_material_score_gauge(events_by_category: Dict[str, List[Dict]]) -> io.BytesIO:
    """
    Create a gauge chart showing average materiality score.

    Args:
        events_by_category: Dict with event data

    Returns:
        BytesIO object containing the chart image
    """
    # Calculate average score across all events
    all_scores = []
    for events in events_by_category.values():
        for event in events:
            if event.get('material_score'):
                all_scores.append(event['material_score'])

    if not all_scores:
        return None

    avg_score = float(sum(all_scores) / len(all_scores))

    # Create gauge chart
    fig, ax = plt.subplots(figsize=(8, 5), subplot_kw={'projection': 'polar'})

    # Gauge parameters
    theta = (avg_score / 10) * 180  # Convert 0-10 to 0-180 degrees

    # Draw gauge arc
    angles = range(0, 181)
    colors_map = plt.cm.RdYlGn([(a/180) for a in angles])

    for i in range(len(angles)-1):
        ax.barh(1, width=1, left=angles[i] * (3.14159/180), height=0.3,
                color=colors_map[i], alpha=0.8)

    # Draw needle
    needle_angle = theta * (3.14159/180)
    ax.plot([needle_angle, needle_angle], [0, 1], color='black', linewidth=3)
    ax.plot(needle_angle, 1, 'o', color='black', markersize=10)

    # Set limits and remove unnecessary elements
    ax.set_ylim(0, 1.3)
    ax.set_xlim(0, 3.14159)  # 0 to π (180 degrees)
    ax.set_xticks([])
    ax.set_yticks([])
    ax.spines['polar'].set_visible(False)

    # Add labels
    ax.text(0, 1.15, 'Low\n(1-3)', ha='center', va='center', fontsize=10)
    ax.text(3.14159/2, 1.15, 'Medium\n(4-6)', ha='center', va='center', fontsize=10)
    ax.text(3.14159, 1.15, 'High\n(7-10)', ha='center', va='center', fontsize=10)

    # Add score text in center
    ax.text(3.14159/2, 0.5, f'{avg_score:.1f}', ha='center', va='center',
            fontsize=36, weight='bold', color='#2E86AB')
    ax.text(3.14159/2, 0.2, 'Average Material Score', ha='center', va='center',
            fontsize=12, style='italic')

    plt.tight_layout()

    # Save to BytesIO
    img_stream = io.BytesIO()
    plt.savefig(img_stream, format='png', dpi=300, bbox_inches='tight')
    img_stream.seek(0)
    plt.close()

    return img_stream


def create_category_comparison_chart(events_by_category: Dict[str, List[Dict]]) -> io.BytesIO:
    """
    Create a grouped bar chart comparing event count and avg material score by category.

    Args:
        events_by_category: Dict with event data

    Returns:
        BytesIO object containing the chart image
    """
    categories = []
    event_counts = []
    avg_scores = []

    for category, events in events_by_category.items():
        if events:
            categories.append(category)
            event_counts.append(len(events))

            scores = [e['material_score'] for e in events if e.get('material_score')]
            avg_scores.append(sum(scores) / len(scores) if scores else 0)

    if not categories:
        return None

    # Create figure with two y-axes
    fig, ax1 = plt.subplots(figsize=(10, 6))
    ax2 = ax1.twinx()

    x = range(len(categories))
    width = 0.35

    # Event counts (bars)
    bars1 = ax1.bar([i - width/2 for i in x], event_counts, width,
                    label='Event Count', color='#2E86AB', alpha=0.8)

    # Avg scores (bars)
    bars2 = ax2.bar([i + width/2 for i in x], avg_scores, width,
                    label='Avg Material Score', color='#F18F01', alpha=0.8)

    # Labels and title
    ax1.set_xlabel('Category', fontsize=12, weight='bold')
    ax1.set_ylabel('Number of Events', fontsize=12, weight='bold', color='#2E86AB')
    ax2.set_ylabel('Average Material Score', fontsize=12, weight='bold', color='#F18F01')
    ax1.set_xticks(x)
    ax1.set_xticklabels(categories)
    ax1.set_title('Event Count and Material Score by Category', fontsize=16, weight='bold', pad=20)

    # Set y2 scale to match 0-10
    ax2.set_ylim(0, 10)

    # Color y-axis labels
    ax1.tick_params(axis='y', labelcolor='#2E86AB')
    ax2.tick_params(axis='y', labelcolor='#F18F01')

    # Add legends
    ax1.legend(loc='upper left')
    ax2.legend(loc='upper right')

    # Add value labels on bars
    for bar in bars1:
        height = bar.get_height()
        ax1.text(bar.get_x() + bar.get_width()/2., height,
                f'{int(height)}', ha='center', va='bottom', fontsize=10, weight='bold')

    for bar in bars2:
        height = bar.get_height()
        ax2.text(bar.get_x() + bar.get_width()/2., height,
                f'{height:.1f}', ha='center', va='bottom', fontsize=10, weight='bold')

    plt.tight_layout()

    # Save to BytesIO
    img_stream = io.BytesIO()
    plt.savefig(img_stream, format='png', dpi=300, bbox_inches='tight')
    img_stream.seek(0)
    plt.close()

    return img_stream


def create_document_coverage_chart(events_by_category: Dict[str, List[Dict]], top_n: int = 15) -> io.BytesIO:
    """
    Create a bar chart showing events sorted by number of underlying documents.

    Args:
        events_by_category: Dict with event data
        top_n: Number of top events to show

    Returns:
        BytesIO object containing the chart image
    """
    # Collect all events with their document counts
    event_data = []
    for category, events in events_by_category.items():
        for event in events:
            doc_count = len(event.get('source_doc_ids', []))
            event_data.append({
                'name': event['event_name'][:40] + '...' if len(event['event_name']) > 40 else event['event_name'],
                'count': doc_count,
                'category': category
            })

    # Sort by count and take top N
    event_data.sort(key=lambda x: x['count'], reverse=True)
    top_events = event_data[:top_n]

    if not top_events:
        return None

    # Prepare data
    names = [e['name'] for e in top_events]
    counts = [e['count'] for e in top_events]
    categories = [e['category'] for e in top_events]

    # Color map
    color_map = {'Economic': '#2E86AB', 'Diplomatic': '#A23B72', 'Social': '#F18F01'}
    colors = [color_map[cat] for cat in categories]

    # Create chart
    fig, ax = plt.subplots(figsize=(10, max(8, len(names) * 0.35)))

    bars = ax.barh(names, counts, color=colors)

    # Add value labels
    for i, (bar, count) in enumerate(zip(bars, counts)):
        ax.text(count + 0.5, i, str(count), va='center', fontsize=9, weight='bold')

    ax.set_xlabel('Number of Source Documents', fontsize=12, weight='bold')
    ax.set_title(f'Top {len(names)} Events by Media Coverage', fontsize=16, weight='bold', pad=20)
    ax.invert_yaxis()
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)

    # Add legend
    legend_patches = [mpatches.Patch(color=color_map[cat], label=cat)
                     for cat in color_map.keys()]
    ax.legend(handles=legend_patches, loc='lower right')

    plt.tight_layout()

    # Save to BytesIO
    img_stream = io.BytesIO()
    plt.savefig(img_stream, format='png', dpi=300, bbox_inches='tight')
    img_stream.seek(0)
    plt.close()

    return img_stream


def create_material_score_histogram(events_by_category: Dict[str, List[Dict]]) -> io.BytesIO:
    """
    Create a histogram showing document count distribution by material score bins.

    Args:
        events_by_category: Dict with event data

    Returns:
        BytesIO object containing the chart image
    """
    # Collect all events with their material scores and doc counts
    event_data = []
    for category, events in events_by_category.items():
        for event in events:
            if event.get('material_score'):
                doc_count = len(event.get('source_doc_ids', []))
                event_data.append({
                    'score': float(event['material_score']),
                    'doc_count': doc_count
                })

    if not event_data:
        return None

    # Create bins 1-10
    bins = list(range(1, 12))  # 1-11 to include 10
    bin_labels = [f'{i}' for i in range(1, 11)]

    # Aggregate document counts by score bin
    bin_doc_counts = [0] * 10

    for event in event_data:
        score = event['score']
        doc_count = event['doc_count']
        # Assign to bin (1.0-1.99 -> bin 0, 2.0-2.99 -> bin 1, etc.)
        bin_idx = min(int(score) - 1, 9)  # Cap at bin 9 (score 10)
        if bin_idx >= 0:
            bin_doc_counts[bin_idx] += doc_count

    # Create histogram
    fig, ax = plt.subplots(figsize=(10, 6))

    bars = ax.bar(bin_labels, bin_doc_counts, color='#2E86AB', alpha=0.8, edgecolor='black', linewidth=1.5)

    # Add value labels on bars
    for bar, count in zip(bars, bin_doc_counts):
        if count > 0:
            ax.text(bar.get_x() + bar.get_width()/2., count,
                   f'{int(count)}', ha='center', va='bottom', fontsize=10, weight='bold')

    ax.set_xlabel('Material Score', fontsize=12, weight='bold')
    ax.set_ylabel('Total Number of Source Documents', fontsize=12, weight='bold')
    ax.set_title('Document Coverage Distribution by Material Score', fontsize=16, weight='bold', pad=20)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.grid(axis='y', alpha=0.3, linestyle='--')

    plt.tight_layout()

    # Save to BytesIO
    img_stream = io.BytesIO()
    plt.savefig(img_stream, format='png', dpi=300, bbox_inches='tight')
    img_stream.seek(0)
    plt.close()

    return img_stream


def create_event_count_histogram(events_by_category: Dict[str, List[Dict]]) -> io.BytesIO:
    """
    Create a histogram showing event count distribution by material score bins.

    Args:
        events_by_category: Dict with event data

    Returns:
        BytesIO object containing the chart image
    """
    # Collect all events with their material scores
    event_scores = []
    for category, events in events_by_category.items():
        for event in events:
            if event.get('material_score'):
                event_scores.append(float(event['material_score']))

    if not event_scores:
        return None

    # Create bins 1-10
    bin_labels = [f'{i}' for i in range(1, 11)]

    # Count events in each score bin
    bin_event_counts = [0] * 10

    for score in event_scores:
        # Assign to bin (1.0-1.99 -> bin 0, 2.0-2.99 -> bin 1, etc.)
        bin_idx = min(int(score) - 1, 9)  # Cap at bin 9 (score 10)
        if bin_idx >= 0:
            bin_event_counts[bin_idx] += 1

    # Create histogram
    fig, ax = plt.subplots(figsize=(10, 6))

    bars = ax.bar(bin_labels, bin_event_counts, color='#A23B72', alpha=0.8, edgecolor='black', linewidth=1.5)

    # Add value labels on bars
    for bar, count in zip(bars, bin_event_counts):
        if count > 0:
            ax.text(bar.get_x() + bar.get_width()/2., count,
                   f'{int(count)}', ha='center', va='bottom', fontsize=10, weight='bold')

    ax.set_xlabel('Material Score', fontsize=12, weight='bold')
    ax.set_ylabel('Number of Events', fontsize=12, weight='bold')
    ax.set_title('Event Count Distribution by Material Score', fontsize=16, weight='bold', pad=20)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.grid(axis='y', alpha=0.3, linestyle='--')

    plt.tight_layout()

    # Save to BytesIO
    img_stream = io.BytesIO()
    plt.savefig(img_stream, format='png', dpi=300, bbox_inches='tight')
    img_stream.seek(0)
    plt.close()

    return img_stream


def add_charts_to_document(doc, events_by_category: Dict[str, List[Dict]], country: str):
    """
    Generate all charts and add them to the document with proper formatting.

    Args:
        doc: python-docx Document object
        events_by_category: Event data organized by category
        country: Country name for context
    """
    doc.add_page_break()
    doc.add_heading('Executive Summary - Visualizations', 1)

    # 1. Category Distribution Pie Chart
    try:
        chart_stream = create_category_pie_chart(events_by_category)
        if chart_stream:
            doc.add_paragraph()
            doc.add_picture(chart_stream, width=Inches(5.5))
            doc.add_paragraph()
    except Exception as e:
        print(f"  [WARNING] Failed to create pie chart: {e}")

    # 2. Material Score Gauge
    try:
        chart_stream = create_material_score_gauge(events_by_category)
        if chart_stream:
            doc.add_paragraph()
            doc.add_picture(chart_stream, width=Inches(5.5))
            doc.add_paragraph()
    except Exception as e:
        print(f"  [WARNING] Failed to create gauge chart: {e}")

    # 3. Category Comparison Chart
    try:
        chart_stream = create_category_comparison_chart(events_by_category)
        if chart_stream:
            doc.add_paragraph()
            doc.add_picture(chart_stream, width=Inches(6))
            doc.add_paragraph()
    except Exception as e:
        print(f"  [WARNING] Failed to create comparison chart: {e}")

    # 4. Document Coverage Chart
    try:
        chart_stream = create_document_coverage_chart(events_by_category, top_n=15)
        if chart_stream:
            doc.add_paragraph()
            doc.add_picture(chart_stream, width=Inches(6))
            doc.add_paragraph()
    except Exception as e:
        print(f"  [WARNING] Failed to create coverage chart: {e}")

    # 5. Top Recipients Chart
    try:
        chart_stream = create_top_recipients_chart(events_by_category, top_n=10)
        if chart_stream:
            doc.add_paragraph()
            doc.add_picture(chart_stream, width=Inches(6))
            doc.add_paragraph()
    except Exception as e:
        print(f"  [WARNING] Failed to create recipients chart: {e}")

    # 6. Material Score Histogram (Document Count)
    try:
        chart_stream = create_material_score_histogram(events_by_category)
        if chart_stream:
            doc.add_paragraph()
            doc.add_picture(chart_stream, width=Inches(6))
            doc.add_paragraph()
    except Exception as e:
        print(f"  [WARNING] Failed to create document count histogram: {e}")

    # 7. Event Count Histogram
    try:
        chart_stream = create_event_count_histogram(events_by_category)
        if chart_stream:
            doc.add_paragraph()
            doc.add_picture(chart_stream, width=Inches(6))
            doc.add_paragraph()
    except Exception as e:
        print(f"  [WARNING] Failed to create event count histogram: {e}")

    print(f"  [CHARTS] Added 7 visualization charts to document")
