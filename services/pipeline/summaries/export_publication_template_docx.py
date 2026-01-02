#!/usr/bin/env python3
"""
Export monthly summary publications as DOCX files matching the template format.

This script generates two versions of publication documents:
1. Summary version: Clean, executive-ready format without doc IDs
2. Reviewer version: Same content with doc_id UUIDs for source traceability

Both documents organize events by soft-power category (Economic, Diplomatic, Social)
and filter events by materiality score to include only significant events.

Usage:
    python export_publication_template_docx.py --country China --month 2025-08
    python export_publication_template_docx.py --influencers --month 2025-09
    python export_publication_template_docx.py --country "United States" --start-date 2025-08-01 --end-date 2025-08-31
"""

import argparse
from pathlib import Path
from datetime import datetime, date
from typing import List, Dict
from collections import defaultdict
import re

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from sqlalchemy import and_, extract
import yaml
import requests

from shared.database.database import get_session
from shared.models.models import EventSummary, EventSourceLink, PeriodType, EventStatus
from services.pipeline.summaries.publication_charts import add_charts_to_document


# Category mapping for soft-power types
CATEGORY_MAPPING = {
    'Economic': ['Economic', 'Infrastructure', 'Trade', 'Investment', 'Development'],
    'Diplomatic': ['Diplomatic', 'Political', 'Military', 'Security'],
    'Social': ['Social', 'Cultural', 'Educational', 'Humanitarian', 'Media']
}


def is_mostly_english(text: str, threshold: float = 0.7) -> bool:
    """Check if text is mostly English (>70% ASCII characters)."""
    if not text:
        return True

    # Count ASCII printable characters (English alphabet, numbers, punctuation)
    ascii_chars = sum(1 for c in text if ord(c) < 128)
    total_chars = len(text)

    return (ascii_chars / total_chars) >= threshold


def translate_to_english_via_api(text: str, api_url: str = None) -> str:
    """
    Translate non-English text to English using the FastAPI proxy.
    Falls back to marking the text if translation fails.
    """
    if not text or is_mostly_english(text):
        return text

    # Try to use the FastAPI material_query endpoint for translation
    if api_url:
        try:
            prompt = f"Translate the following text to English. Provide ONLY the translation, no explanations:\n\n{text}"
            response = requests.post(
                f"{api_url}/material_query",
                json={"prompt": prompt},
                timeout=30
            )
            if response.status_code == 200:
                result = response.json()
                translated = result.get('response', '').strip()
                if translated and is_mostly_english(translated):
                    return translated
        except Exception as e:
            print(f"  [WARNING] Translation failed: {e}")

    # Fallback: mark as non-English
    return f"[Non-English Event Name: {text}]"


def sanitize_event_data(event: Dict, api_url: str = None) -> Dict:
    """Ensure all event text fields are in English."""
    sanitized = event.copy()

    # Translate event name if needed
    if not is_mostly_english(sanitized['event_name']):
        print(f"  [TRANSLATING] Non-English event name detected")
        sanitized['event_name'] = translate_to_english_via_api(sanitized['event_name'], api_url)

    # Overview and Outcome should already be in English from LLM generation
    # but we can check and mark if not
    if sanitized['overview'] and not is_mostly_english(sanitized['overview']):
        print(f"  [WARNING] Non-English overview detected")

    if sanitized['outcome'] and not is_mostly_english(sanitized['outcome']):
        print(f"  [WARNING] Non-English outcome detected")

    return sanitized


def categorize_event(categories_list: List[str]) -> str:
    """Determine the primary soft-power category for an event."""
    if not categories_list:
        return 'Social'  # Default to Social for uncategorized

    # Check each category against mapping
    for sp_category, keywords in CATEGORY_MAPPING.items():
        for cat in categories_list:
            if any(keyword.lower() in cat.lower() for keyword in keywords):
                return sp_category

    # Default to Social if no match
    return 'Social'


def get_monthly_summaries_by_category(
    session,
    country: str,
    start_date: date,
    end_date: date
) -> Dict[str, List[Dict]]:
    """
    Get ALL monthly event summaries grouped by soft-power category.
    No filtering by material score - includes all events for comprehensive analysis.

    Args:
        session: Database session
        country: Initiating country
        start_date: Period start
        end_date: Period end

    Returns:
        Dict with keys 'Economic', 'Diplomatic', 'Social' containing event lists
    """
    # Use raw SQL to query ALL summaries (no material score filter)
    from sqlalchemy import text

    query = text("""
        SELECT id, event_name, narrative_summary, count_by_category, material_score
        FROM event_summaries
        WHERE initiating_country = :country
        AND period_type = 'MONTHLY'
        AND period_start >= :start_date
        AND period_end <= :end_date
        AND status = 'ACTIVE'
        AND is_deleted = FALSE
        ORDER BY id
    """)

    results = session.execute(query, {
        'country': country,
        'start_date': start_date,
        'end_date': end_date
    }).fetchall()

    summaries = []
    for row in results:
        summary_id, event_name, narrative, categories, material_score = row
        summaries.append({
            'id': summary_id,
            'event_name': event_name,
            'narrative': narrative or {},
            'categories': categories or {},
            'material_score': material_score
        })

    # Group events by category
    events_by_category = {
        'Economic': [],
        'Diplomatic': [],
        'Social': []
    }

    for summary in summaries:
        # Extract narrative sections
        narrative = summary['narrative']
        overview = narrative.get('monthly_overview', narrative.get('overview', ''))
        outcome = narrative.get('key_outcomes', narrative.get('outcome', ''))

        # Get source document IDs
        source_links = session.query(EventSourceLink).filter(
            EventSourceLink.event_summary_id == summary['id']
        ).all()
        doc_ids = [link.doc_id for link in source_links]

        # Extract categories - if count_by_category is empty, get from source documents
        categories_list = list(summary['categories'].keys()) if summary['categories'] else []

        # If no categories from summary, extract from source documents
        if not categories_list and doc_ids:
            # Query categories from source documents using raw SQL
            cat_query = text("""
                SELECT DISTINCT category
                FROM categories
                WHERE doc_id = ANY(:doc_ids)
            """)
            doc_categories = session.execute(cat_query, {'doc_ids': doc_ids}).fetchall()
            categories_list = [cat[0] for cat in doc_categories]

        # Determine primary category
        sp_category = categorize_event(categories_list)

        event_data = {
            'event_name': summary['event_name'],
            'overview': overview,
            'outcome': outcome,
            'source_doc_ids': doc_ids,
            'material_score': summary['material_score'],
            'categories': categories_list
        }

        events_by_category[sp_category].append(event_data)

    # Sort events within each category by number of source documents (descending)
    for category in events_by_category:
        events_by_category[category].sort(
            key=lambda x: len(x['source_doc_ids']),
            reverse=True
        )

    return events_by_category


def create_summary_docx(
    country: str,
    start_date: date,
    end_date: date,
    events_by_category: Dict[str, List[Dict]],
    output_path: Path
):
    """Create Summary version (no doc IDs)."""
    doc = Document()

    # Add header banner image at the top
    # Path: services/pipeline/summaries -> root
    banner_path = Path(__file__).resolve().parent.parent.parent.parent / "img" / "OSE_Summary.png"
    if banner_path.exists():
        doc.add_picture(str(banner_path), width=Inches(6.5))
        doc.add_paragraph()  # Add spacing after banner
    else:
        print(f"  [WARNING] Banner image not found at {banner_path}")

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Title
    month_str = start_date.strftime("%B %Y")
    title = doc.add_paragraph(f'{country}-MENA: Chinese Soft-Power Initiatives {month_str}')
    title.runs[0].font.size = Pt(16)
    title.runs[0].bold = True

    # Introduction paragraph
    total_events = sum(len(events) for events in events_by_category.values())
    intro = doc.add_paragraph(
        f'During {month_str}, Near Eastern and North African media covered {country}\'s soft-power '
        f'initiatives throughout the region, highlighting {total_events} significant events across '
        f'economic, diplomatic, and social domains.'
    )

    doc.add_paragraph()  # Spacing

    # Add charts/visualizations before the detailed events
    try:
        add_charts_to_document(doc, events_by_category, country)
    except Exception as e:
        print(f"  [WARNING] Failed to add charts: {e}")

    # Process each category
    for category in ['Economic', 'Diplomatic', 'Social']:
        events = events_by_category[category]

        if not events:
            continue

        # Category header
        heading = doc.add_heading(category.upper(), 1)

        # Category introduction
        intro_text = f'Regional media coverage of {country}\'s {category} soft-power initiatives focused on the following'
        doc.add_paragraph(intro_text)
        doc.add_paragraph()  # Spacing

        # Events in this category
        for event in events:
            # Event name (bold)
            event_para = doc.add_paragraph()
            event_para.add_run(event['event_name']).bold = True

            # Overview
            if event['overview']:
                overview_para = doc.add_paragraph()
                overview_para.add_run('Overview: ').bold = True
                overview_para.add_run(event['overview'])

            # Outcomes
            if event['outcome']:
                outcome_para = doc.add_paragraph()
                outcome_para.add_run('Outcomes: ').bold = True
                outcome_para.add_run(event['outcome'])

            doc.add_paragraph()  # Spacing between events

    # Save document
    doc.save(output_path)
    print(f"[SAVED] Summary: {output_path}")


def create_reviewer_docx(
    country: str,
    start_date: date,
    end_date: date,
    events_by_category: Dict[str, List[Dict]],
    output_path: Path
):
    """Create Reviewer version (with doc IDs)."""
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Title
    month_str = start_date.strftime("%B %Y")
    title = doc.add_paragraph(
        f'{country}-Middle East Economic, Social, Military, and Diplomatic Events {month_str}'
    )
    title.runs[0].font.size = Pt(16)
    title.runs[0].bold = True

    # Introduction paragraph
    period_str = f'{start_date.strftime("%d %B %Y")} to {end_date.strftime("%d %B %Y")}'
    intro = doc.add_paragraph(
        f'During the period from {period_str}, Near Eastern and North African media covered the '
        f'following significant {country} soft-power initiatives.'
    )

    doc.add_paragraph()  # Spacing

    # Process each category
    for category in ['Economic', 'Diplomatic', 'Social']:
        events = events_by_category[category]

        if not events:
            continue

        # Category header
        heading = doc.add_heading(category.upper(), 1)

        # Category introduction
        intro_text = f'Regional media coverage of {country}\'s {category} soft-power initiatives focused on the following'
        doc.add_paragraph(intro_text)
        doc.add_paragraph()  # Spacing

        # Events in this category
        for event in events:
            # Event name (bold)
            event_para = doc.add_paragraph()
            event_para.add_run(event['event_name']).bold = True

            # Overview
            if event['overview']:
                overview_para = doc.add_paragraph()
                overview_para.add_run('Overview: ').bold = True
                overview_para.add_run(event['overview'])

                # Add doc IDs after overview
                for doc_id in event['source_doc_ids']:
                    id_para = doc.add_paragraph(doc_id)
                    id_para.runs[0].font.size = Pt(9)
                    id_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)

            # Outcomes
            if event['outcome']:
                outcome_para = doc.add_paragraph()
                outcome_para.add_run('Outcomes: ').bold = True
                outcome_para.add_run(event['outcome'])

                # Add doc IDs after outcomes
                for doc_id in event['source_doc_ids']:
                    id_para = doc.add_paragraph(doc_id)
                    id_para.runs[0].font.size = Pt(9)
                    id_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)

            doc.add_paragraph()  # Spacing between events

    # Save document
    doc.save(output_path)
    print(f"[SAVED] Reviewer: {output_path}")


def export_publications(
    countries: List[str],
    start_date: date,
    end_date: date,
    output_dir: Path
):
    """Export both Summary and Reviewer DOCX publications for specified countries."""
    output_dir.mkdir(parents=True, exist_ok=True)

    # Get API URL from environment for translation
    import os
    api_url = os.getenv('API_URL', 'http://localhost:8000')

    with get_session() as session:
        for country in countries:
            print(f"\nProcessing {country}...")

            # Get ALL monthly summaries grouped by category (no filtering)
            events_by_category = get_monthly_summaries_by_category(
                session, country, start_date, end_date
            )

            total_events = sum(len(events) for events in events_by_category.values())

            if total_events == 0:
                print(f"  [SKIP] No monthly summaries found")
                continue

            print(f"  Found {total_events} total events (sorted by document count):")
            for category, events in events_by_category.items():
                if events:
                    print(f"    {category}: {len(events)} events")

            # Sanitize events to ensure English-only text
            print(f"  Sanitizing event data for English-only output...")
            for category in events_by_category:
                events_by_category[category] = [
                    sanitize_event_data(event, api_url)
                    for event in events_by_category[category]
                ]

            # Generate filename base
            filename_base = f"{country.replace(' ', '_')}_{start_date.isoformat()}_{end_date.isoformat()}"

            # Create Summary DOCX
            summary_path = output_dir / f"{filename_base}_Summary.docx"
            create_summary_docx(country, start_date, end_date, events_by_category, summary_path)

            # Create Reviewer DOCX
            reviewer_path = output_dir / f"{filename_base}_Reviewer.docx"
            create_reviewer_docx(country, start_date, end_date, events_by_category, reviewer_path)

            print(f"  [COMPLETE] {country}")


def main():
    parser = argparse.ArgumentParser(
        description='Export monthly publications in Summary and Reviewer formats'
    )
    parser.add_argument('--country', type=str, help='Single country to process')
    parser.add_argument('--influencers', action='store_true',
                       help='Process all influencer countries from config')
    parser.add_argument('--month', type=str, help='Month in YYYY-MM format (uses full month)')
    parser.add_argument('--start-date', type=str, help='Start date (YYYY-MM-DD)')
    parser.add_argument('--end-date', type=str, help='End date (YYYY-MM-DD)')
    parser.add_argument('--output-dir', type=str, default='_data/publications',
                       help='Output directory for DOCX files')

    args = parser.parse_args()

    # Determine countries
    countries = []
    if args.influencers:
        with open("shared/config/config.yaml") as f:
            countries = yaml.safe_load(f).get("influencers", [])
    elif args.country:
        countries = [args.country]
    else:
        print("Error: Must specify --country or --influencers")
        return

    # Determine date range
    if args.month:
        # Parse month (YYYY-MM)
        year, month = map(int, args.month.split('-'))
        start_date = date(year, month, 1)
        # Get last day of month
        if month == 12:
            end_date = date(year + 1, 1, 1)
        else:
            end_date = date(year, month + 1, 1)
        from datetime import timedelta
        end_date = end_date - timedelta(days=1)
    elif args.start_date and args.end_date:
        start_date = date.fromisoformat(args.start_date)
        end_date = date.fromisoformat(args.end_date)
    else:
        print("Error: Must specify --month OR --start-date and --end-date")
        return

    output_dir = Path(args.output_dir)

    print("=" * 80)
    print("EXPORTING MONTHLY PUBLICATIONS (SUMMARY + REVIEWER)")
    print("=" * 80)
    print(f"Countries: {', '.join(countries)}")
    print(f"Period: {start_date} to {end_date}")
    print(f"Filter: None (all events included, sorted by document count)")
    print(f"Output: {output_dir}")
    print("=" * 80)
    print()

    export_publications(countries, start_date, end_date, output_dir)

    print()
    print("=" * 80)
    print("EXPORT COMPLETE")
    print("=" * 80)


if __name__ == "__main__":
    main()
