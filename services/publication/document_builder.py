"""
Word document builder for publication generation.

Creates both reviewer and summary versions of publications:
- Reviewer version: Document IDs visible under each section for verification
- Summary version: Clean version with end notes
"""

import os
from typing import List, Dict, Any, Optional
from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


class DocumentBuilder:
    """Builds Word documents from templates with event summaries."""

    def __init__(self, template_path: str):
        """
        Initialize document builder.

        Args:
            template_path: Path to the Word template file
        """
        self.template_path = template_path
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Template not found: {template_path}")

    def build_reviewer_version(
        self,
        output_path: str,
        country: str,
        start_date: date,
        end_date: date,
        title: str,
        events_by_category: Dict[str, List[Dict[str, Any]]],
        categories: List[str],
        image_path: Optional[str] = None
    ) -> str:
        """
        Build the reviewer version with visible document citations.

        Args:
            output_path: Where to save the document
            country: Initiating country
            start_date: Period start date
            end_date: Period end date
            title: Publication title
            events_by_category: Dict mapping category -> list of event dicts
            categories: List of category names in order
            image_path: Optional path to atom.png for hyperlinks

        Returns:
            Path to the created document
        """
        # Load template
        doc = Document(self.template_path)

        # Replace global placeholders
        global_placeholders = {
            '{{country}}': country,
            '{{date}}': self._format_full_date_range(start_date, end_date),
            '{{summary_title}}': title
        }

        self._replace_placeholders_globally(doc, global_placeholders)

        # Category placeholders
        category_placeholders = {
            'Economic': '{{economic_event_section}}',
            'Diplomacy': '{{diplomatic_event_section}}',
            'Social': '{{social_event_section}}',
            'Military': '{{military_event_section}}'
        }

        # Get available styles
        available_styles = {s.name for s in doc.styles}

        # Process each category
        for category in categories:
            if category not in category_placeholders:
                continue

            placeholder = category_placeholders[category]
            events = events_by_category.get(category, [])

            # Find paragraph with placeholder
            for para in doc.paragraphs:
                if placeholder in para.text:
                    para.text = ''  # Clear placeholder
                    last_para = para

                    # Add each event
                    for event in events:
                        # Event heading
                        event_para = self._insert_paragraph_after(
                            last_para,
                            style_name='Heading 3' if 'Heading 3' in available_styles else None
                        )
                        run = event_para.add_run(event['event_name'])
                        run.bold = True
                        last_para = event_para

                        # Overview section
                        overview_para = self._insert_paragraph_after(last_para, style_name='Normal')
                        label_run = overview_para.add_run("Overview: ")
                        label_run.bold = True
                        text_run = overview_para.add_run(event.get('overview', ''))
                        overview_para.add_run(" ")

                        # Add hyperlink image if available
                        if image_path and event.get('overview_hyperlink'):
                            self._add_image_hyperlink(
                                overview_para,
                                image_path,
                                event['overview_hyperlink'],
                                width=0.2
                            )
                        last_para = overview_para

                        # Overview citations (visible in reviewer version)
                        for citation in event.get('overview_citations', []):
                            cite_para = self._insert_paragraph_after(last_para, style_name='Normal')
                            fmt = cite_para.paragraph_format
                            fmt.space_before = Pt(0)
                            fmt.space_after = Pt(0)
                            fmt.line_spacing = 1
                            run = cite_para.add_run(citation)
                            run.font.size = Pt(8)
                            last_para = cite_para

                        # Outcome section
                        outcome_para = self._insert_paragraph_after(last_para, style_name='Normal')
                        label_run = outcome_para.add_run("Outcomes: ")
                        label_run.bold = True
                        text_run = outcome_para.add_run(event.get('outcome', ''))
                        outcome_para.add_run(" ")

                        # Add hyperlink image if available
                        if image_path and event.get('outcome_hyperlink'):
                            self._add_image_hyperlink(
                                outcome_para,
                                image_path,
                                event['outcome_hyperlink'],
                                width=0.2
                            )
                        last_para = outcome_para

                        # Outcome citations (visible in reviewer version)
                        for citation in event.get('outcome_citations', []):
                            cite_para = self._insert_paragraph_after(last_para, style_name='Normal')
                            fmt = cite_para.paragraph_format
                            fmt.space_before = Pt(0)
                            fmt.space_after = Pt(0)
                            fmt.line_spacing = 1
                            run = cite_para.add_run(citation)
                            run.font.size = Pt(8)
                            last_para = cite_para

                    break  # Done with this category

        # Save document
        doc.save(output_path)
        return output_path

    def build_summary_version(
        self,
        output_path: str,
        country: str,
        start_date: date,
        end_date: date,
        title: str,
        events_by_category: Dict[str, List[Dict[str, Any]]],
        categories: List[str]
    ) -> str:
        """
        Build the clean summary version with end notes.

        Args:
            output_path: Where to save the document
            country: Initiating country
            start_date: Period start date
            end_date: Period end date
            title: Publication title
            events_by_category: Dict mapping category -> list of event dicts
            categories: List of category names in order

        Returns:
            Path to the created document
        """
        # Load template
        doc = Document(self.template_path)

        # Replace global placeholders
        global_placeholders = {
            '{{country}}': country,
            '{{date}}': self._format_full_date_range(start_date, end_date),
            '{{summary_title}}': title
        }

        self._replace_placeholders_globally(doc, global_placeholders)

        # Category placeholders
        category_placeholders = {
            'Economic': '{{economic_event_section}}',
            'Diplomacy': '{{diplomatic_event_section}}',
            'Social': '{{social_event_section}}',
            'Military': '{{military_event_section}}'
        }

        # Get available styles
        available_styles = {s.name for s in doc.styles}

        # Collect citations for end notes
        event_citations = {cat: {} for cat in categories}

        # Process each category
        for category in categories:
            if category not in category_placeholders:
                continue

            placeholder = category_placeholders[category]
            events = events_by_category.get(category, [])

            # Find paragraph with placeholder
            for para in doc.paragraphs:
                if placeholder in para.text:
                    para.text = ''  # Clear placeholder
                    last_para = para

                    # Add each event
                    for event in events:
                        event_name = event['event_name']

                        # Event heading
                        event_para = self._insert_paragraph_after(
                            last_para,
                            style_name='Heading 3' if 'Heading 3' in available_styles else None
                        )
                        run = event_para.add_run(event_name)
                        run.bold = True
                        last_para = event_para

                        # Overview section
                        overview_para = self._insert_paragraph_after(last_para, style_name='Normal')
                        label_run = overview_para.add_run("Overview: ")
                        label_run.bold = True
                        text_run = overview_para.add_run(event.get('overview', ''))
                        last_para = overview_para

                        # Outcome section
                        outcome_para = self._insert_paragraph_after(last_para, style_name='Normal')
                        label_run = outcome_para.add_run("Outcomes: ")
                        label_run.bold = True
                        text_run = outcome_para.add_run(event.get('outcome', ''))
                        last_para = outcome_para

                        # Collect citations for end notes
                        citations = []
                        citations.extend(event.get('overview_citations', []))
                        citations.extend(event.get('outcome_citations', []))
                        event_citations[category][event_name] = citations

                    break  # Done with this category

        # Add End Notes section
        self._add_end_notes(doc, event_citations, categories, available_styles)

        # Save document
        doc.save(output_path)
        return output_path

    def _add_end_notes(
        self,
        doc: Document,
        event_citations: Dict[str, Dict[str, List[str]]],
        categories: List[str],
        available_styles: set
    ):
        """Add end notes section with deduplicated citations."""
        # Main heading
        if 'Heading 1' in available_styles:
            hn = doc.add_paragraph(style='Heading 1')
        else:
            hn = doc.add_paragraph()
        hn.add_run("End Notes").bold = True

        for category in categories:
            if category not in event_citations:
                continue

            # Category heading
            if 'Heading 2' in available_styles:
                ch = doc.add_paragraph(style='Heading 2')
            else:
                ch = doc.add_paragraph()
            ch.add_run(category).bold = True

            for event_name, citations in event_citations[category].items():
                if not citations:
                    continue

                # Event sub-heading
                if 'Heading 3' in available_styles:
                    eh = doc.add_paragraph(style='Heading 3')
                else:
                    eh = doc.add_paragraph()
                eh.add_run(event_name).bold = True

                # Deduplicate while preserving order
                seen = []
                for cite in citations:
                    if cite not in seen:
                        seen.append(cite)

                # Write each unique citation
                for cite in seen:
                    p = doc.add_paragraph(style='Normal')
                    fmt = p.paragraph_format
                    fmt.space_before = Pt(0)
                    fmt.space_after = Pt(0)
                    fmt.line_spacing = 1
                    run = p.add_run(cite)
                    run.font.size = Pt(8)

    def _replace_placeholders_globally(self, doc: Document, mapping: Dict[str, str]):
        """Replace placeholders throughout the document."""
        # Replace in main document
        for para in doc.paragraphs:
            self._replace_placeholders_in_para(para, mapping)

        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        self._replace_placeholders_in_para(para, mapping)

        # Replace in headers and footers
        for section in doc.sections:
            for para in section.header.paragraphs:
                self._replace_placeholders_in_para(para, mapping)
            for para in section.footer.paragraphs:
                self._replace_placeholders_in_para(para, mapping)

    @staticmethod
    def _replace_placeholders_in_para(para: Paragraph, mapping: Dict[str, str]):
        """Replace placeholders in a single paragraph."""
        for key, val in mapping.items():
            if key in para.text:
                new_text = para.text.replace(key, val)
                # Clear existing runs
                for run in para.runs:
                    run.text = ""
                # Add new text
                para.add_run(new_text)

    @staticmethod
    def _insert_paragraph_after(existing_para: Paragraph, style_name: Optional[str] = None) -> Paragraph:
        """Insert a new paragraph after an existing one."""
        new_p_elm = OxmlElement('w:p')
        existing_para._p.addnext(new_p_elm)
        new_para = Paragraph(new_p_elm, existing_para._parent)

        if style_name and style_name in {s.name for s in existing_para.part.styles}:
            new_para.style = style_name

        return new_para

    @staticmethod
    def _add_image_hyperlink(
        paragraph: Paragraph,
        image_path: str,
        target_url: str,
        width: float = 0.2
    ):
        """Add a clickable image hyperlink to a paragraph."""
        part = paragraph.part

        # Insert the picture
        pic_run = paragraph.add_run()
        pic_run.add_picture(image_path, width=Inches(width))

        # Build hyperlink
        rId_link = part.relate_to(
            target_url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True
        )
        hlink = OxmlElement('w:hyperlink')
        hlink.set(qn('r:id'), rId_link)

        # Move the picture run into the hyperlink
        run_elm = paragraph._p[-1]
        paragraph._p.remove(run_elm)
        hlink.append(run_elm)
        paragraph._p.append(hlink)

    @staticmethod
    def _format_date_range(start_date: date, end_date: date) -> str:
        """Format date range as 'Month Year' or 'Month Year to Month Year'."""
        start_str = start_date.strftime("%B %Y")
        end_str = end_date.strftime("%B %Y")

        if start_str == end_str:
            return start_str
        else:
            return f"{start_str} to {end_str}"

    @staticmethod
    def _format_full_date_range(start_date: date, end_date: date) -> str:
        """Format full date range with days."""
        start_str = start_date.strftime("%d %B %Y")
        end_str = end_date.strftime("%d %B %Y")
        return f"{start_str} to {end_str}"

    @staticmethod
    def build_hyperlink(doc_ids: List[str], base_url: str = "http://localhost:8501") -> str:
        """
        Build a hyperlink URL for document IDs.

        Args:
            doc_ids: List of document IDs
            base_url: Base URL for the application

        Returns:
            Formatted hyperlink URL
        """
        if not doc_ids:
            return ""

        # Join doc_ids with commas
        doc_ids_str = ",".join(str(doc_id) for doc_id in doc_ids)

        # Build URL (adjust based on your actual routing)
        return f"{base_url}/documents?ids={doc_ids_str}"
